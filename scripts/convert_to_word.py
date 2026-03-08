"""
Converte artigo_pibic_rascunho.md para artigo_pibic.docx com formatação ABNT básica.
Uso: python scripts/convert_to_word.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = os.path.join('docs', 'artigo_pibic_rascunho.md')
OUT_PATH = os.path.join('docs', 'artigo_pibic.docx')
IMG_BASE = os.path.dirname(MD_PATH)   # docs/

doc = Document()

# --- Margens ABNT (3cm esq/sup, 2cm dir/inf) ---
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# --- Estilo padrão ---
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def set_paragraph_format(para, space_before=0, space_after=6, line_spacing=Pt(18)):
    fmt = para.paragraph_format
    fmt.space_before    = Pt(space_before)
    fmt.space_after     = Pt(space_after)
    fmt.line_spacing    = line_spacing
    fmt.first_line_indent = Cm(1.25)

def add_heading(doc, text, level):
    para = doc.add_paragraph()
    para.paragraph_format.space_before   = Pt(12)
    para.paragraph_format.space_after    = Pt(6)
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    run.bold = True
    sizes = {1: 14, 2: 13, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def add_image(doc, img_path, caption):
    """Insere imagem centralizada com legenda."""
    if not os.path.exists(img_path):
        doc.add_paragraph(f'[Imagem não encontrada: {img_path}]')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run()
    run.add_picture(img_path, width=Cm(14))
    # Caption
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    for run in cap.runs:
        run.font.size = Pt(10)
        run.italic = True

def add_table_from_md(doc, lines):
    """Parseia uma tabela markdown e adiciona ao documento."""
    rows = []
    for line in lines:
        if re.match(r'\s*\|[-:| ]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # strip markdown bold
            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            if i == 0:
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()  # espaço após tabela

def apply_inline(para, text):
    """Aplica negrito/itálico inline dentro de um parágrafo."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        # texto antes
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(0).startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(0).startswith('*'):
            run = para.add_run(m.group(3))
            run.italic = True
        else:  # code
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])

# ── Lê o markdown ──────────────────────────────────────────────────────────────
with open(MD_PATH, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Linha horizontal
    if re.match(r'^---+$', line.strip()):
        i += 1
        continue

    # Imagem: ![caption](path)
    img_match = re.match(r'!\[(.+?)\]\((.+?)\)', line.strip())
    if img_match:
        caption  = img_match.group(1)
        rel_path = img_match.group(2)
        abs_path = os.path.normpath(os.path.join(IMG_BASE, rel_path))
        add_image(doc, abs_path, caption)
        i += 1
        continue

    # Tabela markdown
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_table_from_md(doc, table_lines)
        continue

    # Cabeçalhos
    h1 = re.match(r'^# (.+)', line)
    h2 = re.match(r'^## (.+)', line)
    h3 = re.match(r'^### (.+)', line)
    if h1:
        add_heading(doc, h1.group(1), 1)
        i += 1
        continue
    if h2:
        add_heading(doc, h2.group(1), 2)
        i += 1
        continue
    if h3:
        add_heading(doc, h3.group(1), 3)
        i += 1
        continue

    # Metadados título (**Título:** etc.)
    meta_match = re.match(r'^\*\*(.+?):\*\*\s*(.+)', line)
    if meta_match:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(f'{meta_match.group(1)}: ')
        run.bold = True
        para.add_run(meta_match.group(2))
        i += 1
        continue

    # Palavras-chave em itálico (*texto*)
    if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(line.strip().strip('*'))
        run.italic = True
        i += 1
        continue

    # Nota em itálico *(Nota: ...)*
    nota_match = re.match(r'^\*\((.+)\)\.\*$', line.strip())
    if nota_match:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(f'({nota_match.group(1)}).')
        run.italic = True
        run.font.size = Pt(10)
        i += 1
        continue

    # Linha em branco
    if line.strip() == '':
        i += 1
        continue

    # Parágrafo normal
    text = line.strip()
    # remove markdown de nota científica *(Nota ...)*
    text = re.sub(r'^\*\((.+)\)\.\*$', r'(\1).', text)

    para = doc.add_paragraph()
    set_paragraph_format(para)
    apply_inline(para, text)
    i += 1

doc.save(OUT_PATH)
print(f"Documento salvo em: {OUT_PATH}")
